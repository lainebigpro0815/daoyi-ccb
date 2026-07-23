from datetime import date, timedelta
from io import BytesIO
from pathlib import Path
from sqlalchemy.orm import Session
from app.models.project import Project, ProjectPhase, ProjectTask

STATUS_CN = {
    "pending": "待开始",
    "in_progress": "进行中",
    "completed": "已完成",
    "blocked": "阻塞",
}


def _get_week_range() -> tuple[date, date]:
    """获取本周的日期范围（周一 ~ 周日）"""
    today = date.today()
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    return monday, sunday


def _collect_project_data(db: Session, project_id: int) -> dict:
    """收集项目数据，返回结构化字典"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        return {}

    monday, sunday = _get_week_range()
    phases_data = []

    total_tasks = 0
    completed_tasks = 0
    in_progress_tasks = 0
    pending_tasks = 0
    overdue_tasks = 0

    phases = (
        db.query(ProjectPhase)
        .filter(ProjectPhase.project_id == project_id)
        .order_by(ProjectPhase.sort_order)
        .all()
    )

    for phase in phases:
        tasks = (
            db.query(ProjectTask)
            .filter(ProjectTask.project_phase_id == phase.id)
            .order_by(ProjectTask.sort_order)
            .all()
        )

        pc = sum(1 for t in tasks if t.status == "completed")
        ip = sum(1 for t in tasks if t.status == "in_progress")
        pd = sum(1 for t in tasks if t.status == "pending")
        od = sum(
            1
            for t in tasks
            if t.status in ("pending", "in_progress")
            and t.planned_end
            and t.planned_end < date.today()
        )

        total_tasks += len(tasks)
        completed_tasks += pc
        in_progress_tasks += ip
        pending_tasks += pd
        overdue_tasks += od

        phase_tasks = [
            {
                "number": t.task_number,
                "name": t.name,
                "assignee": t.assignee or "未分配",
                "progress": t.progress,
                "status": t.status,
                "planned_end": str(t.planned_end or ""),
                "is_overdue": bool(
                    t.status in ("pending", "in_progress")
                    and t.planned_end
                    and t.planned_end < date.today()
                ),
            }
            for t in tasks
        ]

        phases_data.append(
            {
                "phase_number": phase.phase_number,
                "name": phase.name,
                "status": phase.status,
                "planned_start": str(phase.planned_start or ""),
                "planned_end": str(phase.planned_end or ""),
                "tasks": phase_tasks,
                "task_stats": {
                    "total": len(tasks),
                    "completed": pc,
                    "in_progress": ip,
                    "pending": pd,
                },
            }
        )

    return {
        "project": project,
        "week_range": (monday, sunday),
        "phases": phases_data,
        "stats": {
            "total": total_tasks,
            "completed": completed_tasks,
            "in_progress": in_progress_tasks,
            "pending": pending_tasks,
            "overdue": overdue_tasks,
        },
    }


def _build_ai_context(data: dict) -> str:
    """构建用于 AI 生成周报的项目上下文字符串"""
    project = data["project"]
    monday, sunday = data["week_range"]

    lines = [
        f"项目名称：{project.name}",
        f"客户名称：{project.customer_name or '无'}",
        f"项目阶段：{project.stage}",
        f"周报周期：{monday.isoformat()} ~ {sunday.isoformat()}",
        f"状态：{project.status}",
        "",
        "=== 各阶段进展 ===",
    ]

    for phase in data["phases"]:
        lines.append(
            f"  阶段{phase['phase_number']} [{phase['status']}]: {phase['name']}"
            f"（{phase['planned_start']} ~ {phase['planned_end']}）"
        )
        for t in phase["tasks"]:
            overdue_mark = "【已逾期】" if t["is_overdue"] else ""
            lines.append(
                f"    - {t['number']} {t['name']}"
                f"（负责人: {t['assignee']}, 进度: {t['progress']}%,"
                f" 状态: {STATUS_CN.get(t['status'], t['status'])}）{overdue_mark}"
            )
        lines.append("")

    stats = data["stats"]
    lines.extend([
        "=== 统计 ===",
        f"任务总数：{stats['total']}",
        f"已完成：{stats['completed']}",
        f"进行中：{stats['in_progress']}",
        f"待开始：{stats['pending']}",
        f"已逾期：{stats['overdue']}",
        "",
        f"完成率：{round(stats['completed'] / stats['total'] * 100) if stats['total'] else 0}%",
    ])

    return "\n".join(lines)


def _generate_template_report(data: dict) -> str:
    """基于模板生成周报（AI 不可用时的降级方案）"""
    project = data["project"]
    monday, sunday = data["week_range"]
    stats = data["stats"]
    progress_pct = round(stats["completed"] / stats["total"] * 100) if stats["total"] else 0

    report = []
    report.append("项目周报")
    report.append("")
    report.append(f"项目名称：{project.name}")
    report.append(f"客户名称：{project.customer_name or '无'}")
    report.append(f"周报周期：{monday.isoformat()} ~ {sunday.isoformat()}")
    report.append(f"项目阶段：{project.stage}")
    report.append("")

    report.append("一、本周进展总结")
    report.append("")
    report.append(
        f"本周项目处于「运行中」状态。"
        f"已完成 {stats['completed']}/{stats['total']} 个任务（{progress_pct}%），"
        f"其中进行中 {stats['in_progress']} 个，待开始 {stats['pending']} 个"
        f"{'，已逾期 ' + str(stats['overdue']) + ' 个' if stats['overdue'] else ''}。"
    )
    report.append("")

    report.append("二、各阶段详情")
    report.append("")
    for phase in data["phases"]:
        report.append(f"阶段{phase['phase_number']}：{phase['name']}")
        report.append(f"状态：{STATUS_CN.get(phase['status'], phase['status'])}")
        report.append(f"计划周期：{phase['planned_start']} ~ {phase['planned_end']}")
        report.append("")
        if phase["tasks"]:
            for t in phase["tasks"]:
                overdue = " **已逾期**" if t["is_overdue"] else ""
                report.append(
                    f"  [{STATUS_CN.get(t['status'], t['status'])}] "
                    f"{t['number']} {t['name']}"
                    f"（负责人: {t['assignee']}, 进度: {t['progress']}%）{overdue}"
                )
        else:
            report.append("  （本阶段暂无任务）")
        report.append("")

    report.append("三、下周计划")
    report.append("")
    report.append(
        f"下周将重点推进 {stats['in_progress']} 个进行中任务，"
        f"并启动 {stats['pending']} 个待开始任务。"
    )
    if stats["overdue"]:
        report.append(f"其中 {stats['overdue']} 个任务已逾期，需优先安排处理。")
    report.append("")

    report.append("四、风险与问题")
    report.append("")
    if stats["overdue"]:
        report.append(f"- 存在 {stats['overdue']} 个逾期任务，可能影响整体进度，需关注。")
    else:
        report.append("- 暂无重大风险。")
    report.append("")

    report.append("---")
    report.append(f"报告生成日期：{date.today().isoformat()}")

    return "\n".join(report)


async def generate_weekly_report(db: Session, project_id: int) -> str:
    """生成项目周报，优先使用 AI，不可用时降级为模板"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError("项目不存在")

    data = _collect_project_data(db, project_id)

    # 尝试 AI 生成
    from app.services.ai_service import _load_config, stream_ai_response

    cfg = _load_config()
    if cfg.get("provider") != "mock" and cfg.get("api_key"):
        context = _build_ai_context(data)
        prompt = (
            "请根据以下项目数据生成一份专业的项目周报。\n"
            "周报应包含：\n"
            "1）本周进展总结（总体进度、完成情况）\n"
            "2）各阶段详细进展（分阶段列出任务状态）\n"
            "3）风险和问题（逾期任务、潜在风险）\n"
            "4）下周工作计划（重点工作安排）\n"
            "请用中文撰写，使用 Markdown 格式。\n\n"
            f"{context}"
        )
        try:
            full_text = ""
            async for chunk in stream_ai_response(
                context, [{"role": "user", "content": prompt}]
            ):
                full_text += chunk
            if full_text.strip():
                return full_text.strip()
        except Exception:
            pass

    # 降级为模板
    return _generate_template_report(data)


def generate_report_docx(report_text: str, project_name: str) -> bytes:
    """将报告文本转为 Word 文档并返回字节"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    title = doc.add_heading(project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("项 目 周 报")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # 解析 markdown 风格文本并写入 Word
    for line in report_text.split("\n"):
        stripped = line.strip()

        if stripped.startswith("# ") and not stripped.startswith("## "):
            doc.add_heading(stripped[2:].strip(), level=1)

        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)

        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)

        elif stripped.startswith("|"):
            if not all(c in "-| " for c in stripped):
                cells = [c.strip() for c in stripped.split("|") if c.strip()]
                if cells:
                    p = doc.add_paragraph(" | ".join(cells))
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 60)

        elif stripped == "":
            doc.add_paragraph()

        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
